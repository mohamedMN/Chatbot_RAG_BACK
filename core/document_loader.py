"""
Module de chargement de documents de différents formats
"""
import json
import re
import sys
import tempfile
import subprocess
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from utils.file_utils import load_json
from config.settings import SUPPORTED_EXTENSIONS

# Import conditionnel des bibliothèques optionnelles
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
    USE_PYMUPDF = False
except ImportError:
    try:
        import fitz
        PDF_AVAILABLE = True
        USE_PYMUPDF = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PYMUPDF = False


class DocumentLoader:
    """Chargeur de documents multi-formats"""

    def __init__(self):
        self.supported_extensions = SUPPORTED_EXTENSIONS

    def load_document(self, filepath: Path) -> Dict[str, Any]:
        """
        Charge un document depuis un fichier
        
        Args:
            filepath: Chemin du fichier
            
        Returns:
            Document avec contenu et métadonnées
        """
        doc_id = filepath.stem
        file_extension = filepath.suffix.lower()

        try:
            if file_extension == '.txt':
                content, sections = self._load_txt(filepath)
            elif file_extension == '.json':
                content, sections = self._load_json(filepath, doc_id)
            elif file_extension == '.docx':
                content, sections = self._load_docx(filepath)
            elif file_extension == '.pdf':
                content, sections = self._load_pdf(filepath)
            elif file_extension == '.doc':
                content, sections = self._load_doc(filepath)
            else:
                raise ValueError(f"Format non supporté: {file_extension}")

            return {
                'id': doc_id,
                'content': content,
                'source': str(filepath),
                'predefined_sections': sections
            }

        except Exception as e:
            return {
                'id': doc_id,
                'content': f"ERREUR lors du chargement: {str(e)}",
                'source': str(filepath),
                'predefined_sections': []
            }

    def load_documents(self, directory: Path) -> List[Dict[str, Any]]:
        """
        Charge tous les documents d'un dossier
        
        Args:
            directory: Dossier contenant les documents
            
        Returns:
            Liste des documents chargés
        """
        documents = []

        if not directory.exists():
            directory.mkdir(parents=True, exist_ok=True)
            print(f"Dossier créé: {directory}")
            print(f"Veuillez placer vos documents dans: {directory}")
            return []

        files_found = False
        for filepath in directory.iterdir():
            if filepath.is_file() and filepath.suffix.lower() in self.supported_extensions:
                document = self.load_document(filepath)
                documents.append(document)
                files_found = True
                print(f"Document chargé: {filepath.name}")

        if not files_found:
            print(f"Aucun document supporté trouvé dans {directory}")
            print(f"Formats supportés: {', '.join(self.supported_extensions)}")

        return documents

    def _load_txt(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """Charge un fichier texte"""
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content, []  # Pas de sections prédéfinies pour les fichiers texte

    def _load_json(self, filepath: Path, doc_id: str) -> Tuple[str, List[Dict[str, Any]]]:
        """Charge un fichier JSON"""
        doc_json = load_json(filepath)

        if isinstance(doc_json, dict) and 'content' in doc_json:
            content = doc_json['content']
            if 'id' in doc_json:
                doc_id = doc_json['id']
        else:
            content = json.dumps(doc_json, ensure_ascii=False, indent=2)

        return content, []  # Pas de sections prédéfinies pour les fichiers JSON

    def _load_docx(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Charge un fichier DOCX avec extraction des sections
        
        Returns:
            Tuple (contenu, sections)
        """
        if not DOCX_AVAILABLE:
            return "ERREUR: bibliothèque python-docx non installée", []

        try:
            doc = docx.Document(filepath)
            full_text = []
            sections = []
            char_position = 0

            # Extrait le texte et détecte les titres
            for para in doc.paragraphs:
                para_text = para.text.strip()

                if not para_text:
                    continue

                # Vérifie si le paragraphe est un titre
                is_heading = self._is_docx_heading(para)

                if is_heading:
                    sections.append({
                        'title': para_text,
                        'start': char_position
                    })

                full_text.append(para_text)
                char_position += len(para_text) + 1

            # Ajoute le contenu des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            full_text.append(cell_text)
                            char_position += len(cell_text) + 1

            # Finalise les plages des sections
            self._finalize_section_ranges(sections, char_position)

            content = '\n'.join(full_text)
            return content, sections

        except Exception as e:
            return f"ERREUR lecture DOCX: {str(e)}", []

    def _is_docx_heading(self, para) -> bool:
        """Détermine si un paragraphe DOCX est un titre"""
        # Vérifie le style
        if para.style.name.startswith('Heading') or 'Title' in para.style.name:
            return True

        # Vérifie le formatage (gras, police plus grande)
        if any(run.bold for run in para.runs):
            return True

        # Vérifie les patterns courants de titre
        text = para.text.strip()
        if (re.match(r'^[A-Z0-9][\.\)][^\n]+$', text) or
            re.match(r'^[A-Z][A-Za-z\s]{0,20}:$', text) or
            (text.isupper() and len(text) < 80) or
            re.match(r'^Environnement\s+[A-Z]+', text) or
                re.search(r'description\s+technique', text.lower())):
            return True

        return False

    def _load_pdf(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Charge un fichier PDF avec extraction des sections
        
        Returns:
            Tuple (contenu, sections)
        """
        if not PDF_AVAILABLE:
            return "ERREUR: bibliothèque PDF non installée", []

        try:
            sections = []

            if USE_PYMUPDF:
                content, sections = self._load_pdf_pymupdf(filepath)
            else:
                content, sections = self._load_pdf_pypdf2(filepath)

            return content, sections

        except Exception as e:
            return f"ERREUR lecture PDF: {str(e)}", []

    def _load_pdf_pymupdf(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """Charge un PDF avec PyMuPDF (meilleure structure)"""
        import fitz

        doc = fitz.open(filepath)
        full_text = []
        sections = []
        char_position = 0

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()

            # Détecte les titres par taille de police et format
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        if "spans" in line:
                            for span in line["spans"]:
                                text = span.get("text", "").strip()
                                font_size = span.get("size", 0)
                                font_flags = span.get("flags", 0)

                                if self._is_pdf_heading(text, font_size, font_flags):
                                    sections.append({
                                        'title': text,
                                        'start': char_position
                                    })

            full_text.append(page_text)
            char_position += len(page_text)

        self._finalize_section_ranges(sections, char_position)
        content = "\n".join(full_text)

        return content, sections

    def _load_pdf_pypdf2(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """Charge un PDF avec PyPDF2 (moins d'informations structurelles)"""
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = []
            sections = []
            char_position = 0

            for page in reader.pages:
                page_text = page.extract_text()

                # Analyse les lignes pour détecter les titres
                lines = page_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line and self._is_text_heading(line):
                        sections.append({
                            'title': line,
                            'start': char_position
                        })
                    char_position += len(line) + 1

                full_text.append(page_text)

            self._finalize_section_ranges(sections, char_position)
            content = "\n".join(full_text)

            return content, sections

    def _is_pdf_heading(self, text: str, font_size: float, font_flags: int) -> bool:
        """Détermine si un texte PDF est un titre"""
        if not text or len(text) > 100:
            return False

        # Police plus grande que la normale
        if font_size > 12:
            return True

        # Texte en gras (flag 16)
        if font_flags & 16:
            return True

        # Patterns de titre courants
        return self._is_text_heading(text)

    def _is_text_heading(self, text: str) -> bool:
        """Détermine si un texte est un titre (patterns génériques)"""
        if len(text) >= 80:
            return False

        return (text.isupper() or
                re.match(r'^[0-9]+\.\s+', text) or
                re.search(r'description\s+technique', text.lower()) or
                re.match(r'^Environnement\s+[A-Z]+', text) or
                re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){1,3}:', text))

    def _load_doc(self, filepath: Path) -> Tuple[str, List[Dict[str, Any]]]:
        """Charge un fichier DOC (ancien format Word) sans textract.

        Stratégie:
        1) Windows + Word installé -> COM (win32com) -> export TXT
        2) LibreOffice (soffice) présent -> convert-to txt:Text
        3) Échec -> message d'erreur explicite
        """
        try:
            if not filepath.exists() or filepath.suffix.lower() != ".doc":
                return "ERREUR: Fichier manquant ou extension non .doc", []

            # 1) Fallback Word COM (Windows uniquement)
            if sys.platform.startswith("win"):
                try:
                    import win32com.client  # pywin32
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    with tempfile.TemporaryDirectory() as tmpdir:
                        out_txt = Path(tmpdir) / (filepath.stem + ".txt")
                        doc = word.Documents.Open(str(filepath))
                        # 2 = wdFormatText
                        doc.SaveAs(str(out_txt), FileFormat=2)
                        doc.Close(False)
                        word.Quit()
                        content = out_txt.read_text(
                            encoding="utf-8", errors="replace")
                        return content, []
                except Exception:
                    # si Word/pywin32 indisponible, on tente LibreOffice
                    pass

            # 2) Fallback LibreOffice (toutes plateformes) si 'soffice' disponible
            try:
                # Test rapide de présence
                probe = subprocess.run(
                    ["soffice", "--version"], capture_output=True, text=True)
                if probe.returncode == 0:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        tmpdir_p = Path(tmpdir)
                        # Convertir en texte brut
                        # (sur certaines installations: 'soffice.exe' ou 'soffice')
                        cmd = [
                            "soffice",
                            "--headless",
                            "--convert-to", "txt:Text",
                            "--outdir", str(tmpdir_p),
                            str(filepath),
                        ]
                        r = subprocess.run(cmd, capture_output=True, text=True)
                        if r.returncode == 0:
                            out_txt = tmpdir_p / (filepath.stem + ".txt")
                            if out_txt.exists():
                                content = out_txt.read_text(
                                    encoding="utf-8", errors="replace")
                                return content, []
            except FileNotFoundError:
                # soffice introuvable
                pass
            except Exception:
                pass

            # 3) Échec total
            return (
                "ERREUR: Impossible de lire le fichier DOC. "
                "Installe Microsoft Word (avec pywin32) ou LibreOffice (soffice) "
                "ou convertis le fichier en DOCX puis relance.",
                []
            )

        except Exception as e:
            return f"ERREUR lecture DOC: {e}", []

    def _finalize_section_ranges(self, sections: List[Dict[str, Any]], total_length: int) -> None:
        """
        Finalise les plages des sections
        
        Args:
            sections: Liste des sections à finaliser
            total_length: Longueur totale du texte
        """
        # Assigne les positions de fin
        for i in range(len(sections) - 1):
            sections[i]['end'] = sections[i+1]['start'] - 1

        # Assigne la fin pour la dernière section
        if sections:
            sections[-1]['end'] = total_length
