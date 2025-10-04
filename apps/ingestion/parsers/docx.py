# backend/apps/ingestion/parsers/docx.py
from io import BytesIO
from typing import List, Optional
import re
from docx import Document
from docx.text.paragraph import Paragraph
from apps.ingestion.schema import Section

_NUM_HEADING = re.compile(r"^\s*\d+(\.\d+)*\s+[^\n]{1,120}$")


def _name_chain(p: Paragraph) -> list[str]:
    names: list[str] = []
    try:
        st = p.style
        while st is not None:
            n = (st.name or "").lower()
            if n and n not in names:
                names.append(n)
            st = getattr(st, "base_style", None)
    except Exception:
        pass
    return names


def _is_heading_by_style(p: Paragraph) -> bool:
    names = _name_chain(p)
    return any(
        ("heading" in n) or ("titre" in n) or ("title" in n)
        for n in names
    )


def _is_heading_by_heuristics(p: Paragraph, txt: str) -> bool:
    # short, bold/large lines → likely headings
    if len(txt) <= 120:
        bold = any(getattr(r.font, "bold", False)
                   for r in p.runs if r.text.strip())
        # large font (~>= 14pt) counts as a hint
        large = any(
            (getattr(r.font, "size", None) or 0) and (
                r.font.size.pt if r.font.size else 0) >= 14
            for r in p.runs if r.text.strip()
        )
        if bold or large:
            return True
        # numbered heading patterns: "1.", "1.1.", etc.
        if _NUM_HEADING.match(txt):
            return True
        # very uppercase-ish short lines
        letters = [c for c in txt if c.isalpha()]
        if letters:
            up_ratio = sum(c.isupper() for c in letters) / len(letters)
            if up_ratio >= 0.6:
                return True
    return False


def _is_heading(p: Paragraph) -> bool:
    txt = (p.text or "").strip()
    if not txt:
        return False
    return _is_heading_by_style(p) or _is_heading_by_heuristics(p, txt)


def parse_docx_bytes(data: bytes, filename: str) -> List[Section]:
    doc = Document(BytesIO(data))

    sections: List[Section] = []
    current_title: Optional[str] = None
    current_lines: List[str] = []
    ordinal = 0

    # If the file has a core title, use it as the first section title
    doc_title = (getattr(doc.core_properties, "title", "") or "").strip()
    if doc_title and doc_title.lower() not in {"title", "document"}:
        current_title = doc_title[:120]

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue

        if _is_heading(p):
            # Flush previous section
            if (current_title or current_lines):
                sections.append(
                    Section(
                        subject=current_title or "Section",
                        source=f"{filename}#h={current_title or ordinal}",
                        content="\n".join(current_lines).strip(),
                        ordinal=ordinal,
                    )
                )
                ordinal += 1
                current_lines = []
            # Start a new section with this heading
            current_title = txt[:120]
        else:
            current_lines.append(txt)

    # Flush tail
    if (current_title or current_lines):
        sections.append(
            Section(
                subject=current_title or "Section",
                source=f"{filename}#h={current_title or ordinal}",
                content="\n".join(current_lines).strip(),
                ordinal=ordinal,
            )
        )

    # Fallback: whole document as one section
    if not sections:
        text = "\n".join(
            p.text for p in doc.paragraphs if (p.text or "").strip())
        sections = [Section(subject=doc_title or "Document",
                            source=filename, content=text, ordinal=0)]

    return sections
